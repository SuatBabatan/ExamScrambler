from docx import Document
import random
import string
import traceback
from copy import deepcopy
try:
    print('ExamScrambler by Suat Babatan\nhttps://github.com/SuatBabatan')
    alphabet = list(string.ascii_lowercase)
    print('')
    print('make sure you have your questions in questions.docx and answers in answers.docx')
    print('')
    n_choices = int(input('How many choices is there for each question?(enter a number): '))
    n_booklets = int(input('How many shuffled exams do you need?(enter a number): '))


    questions={}
    cheat_sheet = {}

    exam: Document = Document('questions.docx')
    answer_doc: Document = Document('answers.docx')


    i=0
    q_no = 0
    for p in exam.paragraphs:

        if i % (n_choices + 1) == 0:
            key =answer_doc.paragraphs[q_no].text.lower()
            correct_answer_no = alphabet.index(key) % (n_choices + 1) +1
            q_no+=1
            questions[p.text] = set()
            question = p.text
        else:
            if i % (n_choices + 1) == correct_answer_no:
                cheat_sheet[question] = p.text
            ans = p.text
            questions[question].add(ans)
        i+=1

    for i in range(1,n_booklets+1):
        question_booklet: Document = Document()
        answer_booklet: Document = Document()
        questions_tmp = deepcopy(questions)
        for j in range(len(questions_tmp)):
            question_str = random.choice(list(questions_tmp.keys()))
            question = questions_tmp.pop(question_str)
            question_booklet.add_paragraph(question_str, style='List Number')
            key = cheat_sheet[question_str]
            for n in range(len(question)):
                answers = list(question)
                answer = random.choice(answers)
                question.remove(answer)
                if key == answer:
                    answer_booklet.add_paragraph(alphabet[n].upper(), style='List Number')
                question_booklet.add_paragraph(f'                {alphabet[n]}. {answer}')

        question_booklet.save(f'question_booklet{i}.docx')
        answer_booklet.save(f'answer_booklet{i}.docx')
        print('Done! :)')
        print('press any key to exit')
except BaseException as e:
    print('')
    print('')
    traceback.print_exception(type(e), e, e.__traceback__)
    print('')
    print('Woops, we encountered an error :(')
    print('press any key to exit')
    input()